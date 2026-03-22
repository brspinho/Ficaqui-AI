import streamlit as st
import pandas as pd
from fpdf import FPDF
import io
import os
from groq import Groq

def ensure_groq_diagnosis(df, groq_key):
    if not groq_key:
        return "Configuracao da API Groq ausente. O relatorio nao pode gerar insights de IA."
        
    try:
        client = Groq(api_key=groq_key)
        
        df_sample = df.sample(min(15, len(df)))
        context = df_sample.to_json(orient='records')
        
        prompt = f"""Atue como um Consultor de Inteligencia Urbana focado em B2G.
O cliente gerou uma exportacao de dados baseada em filtros especificos. 
Aqui esta uma amostra dos dados filtrados:
{context}

Considerando as estatisticas desse recorte, crie uma "Sintese Executiva" de 2 paragrafos. 
O primeiro deve diagnosticar o cenario (Ex: vacancia, oportunidades). 
O segundo deve propor medidas fiscais/urbanas claras. Nao use ingles."""

        completion = client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.5,
            max_tokens=400
        )
        return completion.choices[0].message.content
    except Exception as e:
        return f"Erro ao contatar a Groq: {e}"

def create_pdf(df, ai_diagnosis):
    import os
    import unicodedata
    from datetime import datetime

    def strip_accents(s):
        return "".join(c for c in unicodedata.normalize('NFD', s) if unicodedata.category(c) != 'Mn')

    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=20)
    
    # --- HEADER / CABECALHO ---
    w1, w2, w3 = 60, 80, 50 
    h_box = 30
    
    logo_path = "assets/logo.png"
    if not os.path.exists(logo_path):
        if os.path.exists("assets"):
            imgs = [f for f in os.listdir("assets") if f.lower().endswith(('.png', '.jpg', '.jpeg'))]
            if imgs:
                logo_path = os.path.join("assets", imgs[0])

    pdf.set_draw_color(0, 0, 0)
    pdf.set_line_width(0.5)
    
    pdf.rect(10, 10, w1, h_box)
    if os.path.exists(logo_path):
        pdf.image(logo_path, x=15, y=12, h=45) 
        
    pdf.rect(10 + w1, 10, w2, h_box)
    pdf.set_xy(10 + w1, 15)
    pdf.set_font("helvetica", 'B', 16)
    pdf.multi_cell(w2, 10, txt=strip_accents("Relatorio Estrategico de\nExportacao de Dados"), border=0, align='C')
    
    pdf.rect(10 + w1 + w2, 10, w3, h_box)
    pdf.set_xy(10 + w1 + w2, 12)
    pdf.set_font("helvetica", 'B', 9)
    pdf.cell(w3, 5, txt=strip_accents("No Relatorio"), ln=True, align='C')
    pdf.set_font("helvetica", '', 8)
    report_id = datetime.now().strftime("%Y%m%d%H%M%S")
    pdf.set_x(10 + w1 + w2)
    pdf.cell(w3, 5, txt=report_id, ln=True, align='C')
    
    pdf.set_xy(10 + w1 + w2, 22)
    pdf.set_font("helvetica", 'B', 9)
    pdf.cell(w3, 5, txt="Data", ln=True, align='C')
    pdf.set_font("helvetica", '', 8)
    pdf.set_x(10 + w1 + w2)
    pdf.cell(w3, 5, txt=datetime.now().strftime("%d/%m/%Y"), ln=True, align='C')
    
    pdf.set_xy(10, 10 + h_box)
    pdf.set_font("helvetica", 'B', 8)
    pdf.set_text_color(50, 50, 50)
    params_txt = f" Parametros: Filtro de Imoveis ({len(df)} ativos) | Engine: Ficaqui AI Llama 3.3 | Analise: Urbanism Strategy"
    pdf.cell(190, 8, txt=strip_accents(params_txt), border=1, ln=True)
    
    pdf.set_xy(10, 55) 
    pdf.set_text_color(0, 0, 0)
    
    # --- ESTISTICAS GERAIS ---
    pdf.set_fill_color(0, 121, 107)
    pdf.set_text_color(255, 255, 255)
    pdf.set_font("helvetica", 'B', 12)
    pdf.cell(0, 10, txt=" SITUACAO DO RECORTE SELECIONADO", ln=True, fill=True)
    pdf.ln(5)
    
    pdf.set_text_color(0, 0, 0)
    pdf.set_font("helvetica", '', 11)
    pdf.cell(0, 8, txt=f"Total de Ativos Analisados: {len(df)} unidades", ln=True)
    receita_total = df['receita_gerada'].sum()
    pdf.cell(0, 8, txt=f"Receita Operacional Mensal Estimada: R$ {receita_total:,.2f}", ln=True)
    pdf.ln(10)
    
    pdf.set_font("helvetica", 'B', 8)
    pdf.set_fill_color(240, 240, 240)
    
    pdf.cell(50, 10, 'Logradouro', 1, 0, 'C', True)
    pdf.cell(30, 10, 'Tipo', 1, 0, 'C', True)
    pdf.cell(40, 10, 'Status', 1, 0, 'C', True)
    pdf.cell(35, 10, 'Fluxo/Dia', 1, 0, 'C', True)
    pdf.cell(35, 10, 'Iluminacao', 1, 1, 'C', True)
    
    pdf.set_font("helvetica", '', 7)
    for i, row in df.head(25).iterrows():
        rua_val = strip_accents(str(row['rua'])[:30]) 
        tipo_val = strip_accents(str(row['tipo']))
        status_val = strip_accents(str(row['status_aluguel']))
        fluxo_val = f"{row['fluxo_pessoas_dia']} hab"
        ilum_val = strip_accents(str(row['iluminacao']))
        
        pdf.cell(50, 8, rua_val, 1, 0, 'L')
        pdf.cell(30, 8, tipo_val, 1, 0, 'C')
        pdf.cell(40, 8, status_val, 1, 0, 'L')
        pdf.cell(35, 8, fluxo_val, 1, 0, 'C')
        pdf.cell(35, 8, ilum_val, 1, 1, 'C')
    
    pdf.ln(15)
    
    pdf.set_font("helvetica", 'B', 14)
    pdf.set_text_color(0, 77, 64)
    pdf.cell(0, 10, txt="Parecer Estrategico (AI Urbanist Analysis)", ln=True)
    pdf.set_draw_color(0, 77, 64)
    pdf.line(10, pdf.get_y(), 200, pdf.get_y())
    pdf.ln(5)
    
    pdf.set_font("helvetica", '', 12)
    pdf.set_text_color(40, 40, 40)
    
    safe_ai_diagnosis = strip_accents(ai_diagnosis)
    pdf.multi_cell(0, 8, safe_ai_diagnosis, align='J', border=0)
    
    pdf.ln(20)
    pdf.set_font("helvetica", 'I', 8)
    pdf.set_text_color(150, 150, 150)
    pdf.cell(0, 10, txt="Plataforma Ficaqui | Aracaju, SE | Inteligencia Artificial Groq/Llama 3.3", align='C')

    return bytes(pdf.output())

def create_docx(df, ai_diagnosis):
    from docx import Document
    from docx.shared import Pt
    import io
    from datetime import datetime

    doc = Document()
    doc.add_heading('Relatorio Estrategico de Escritorio - Ficaqui', 0)

    doc.add_paragraph(f"Data de Emissao: {datetime.now().strftime('%d/%m/%Y')}")
    doc.add_paragraph(f"Total de Ativos Analisados: {len(df)} unidades")
    receita_total = df['receita_gerada'].sum()
    doc.add_paragraph(f"Receita Operacional Mensal Estimada: R$ {receita_total:,.2f}")

    doc.add_heading('Amostra Sintetica do Recorte', level=1)
    
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Light Grid Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Logradouro'
    hdr_cells[1].text = 'Tipo'
    hdr_cells[2].text = 'Status'
    hdr_cells[3].text = 'Fluxo/Dia'
    hdr_cells[4].text = 'Iluminacao'

    for i, row in df.head(25).iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = str(row['rua'])[:30]
        row_cells[1].text = str(row['tipo'])
        row_cells[2].text = str(row['status_aluguel'])
        row_cells[3].text = f"{row['fluxo_pessoas_dia']} hab"
        row_cells[4].text = str(row['iluminacao'])

    doc.add_page_break()
    doc.add_heading('Parecer Estrategico (AI Urbanist Analysis)', level=1)
    
    p = doc.add_paragraph()
    run = p.add_run(ai_diagnosis)
    run.font.size = Pt(12)

    target = io.BytesIO()
    doc.save(target)
    return target.getvalue()

def render_report_view(df, groq_key=None):
    if not groq_key:
        import os
        groq_key = os.getenv("GROQ_API_KEY")

    st.subheader("Relatorios Executivos e Exportacao")
    st.markdown("Filtre a base de dados espaciais e exija que a IA crie pareceres urbanisticos para download oficial.")
    
    col1, col2 = st.columns(2)
    with col1:
        f_status = st.multiselect("Status do Ativo", df['status_aluguel'].unique(), default=df['status_aluguel'].unique())
    with col2:
        f_tipo = st.multiselect("Tipologia de Imovel", df['tipo'].unique(), default=df['tipo'].unique())
        
    df_filtrado = df[(df['status_aluguel'].isin(f_status)) & (df['tipo'].isin(f_tipo))]
    
    st.markdown(f"**Registros Encontrados: {len(df_filtrado)}**")
    st.dataframe(df_filtrado, width='stretch')
    
    c1, c2 = st.columns(2)
    
    with c1:
        csv = df_filtrado.to_csv(index=False).encode('utf-8')
        st.download_button(
            label="Baixar Base Filtrada (CSV)",
            data=csv,
            file_name='ficaqui_export.csv',
            mime='text/csv',
        )
        
    with c2:
        export_format = st.radio("Escolha o Formato do Parecer:", ["PDF", "DOCX (Word)"], horizontal=True)
        
        if st.button("Gerar Parecer Executivo (Groq IA)"):
            with st.spinner("Compilando dados e consultando o modelo Llama 3.3..."):
                ai_diagnosis = ensure_groq_diagnosis(df_filtrado, groq_key)
                st.info(ai_diagnosis)
                
                if export_format == "PDF":
                    file_bytes = create_pdf(df_filtrado, ai_diagnosis)
                    mime_type = "application/pdf"
                    f_name = "ficaqui_executive_report.pdf"
                    f_label = "Baixar Parecer Executivo (PDF)"
                else:
                    file_bytes = create_docx(df_filtrado, ai_diagnosis)
                    mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    f_name = "ficaqui_executive_report.docx"
                    f_label = "Baixar Parecer Executivo (DOCX)"
                    
                st.download_button(
                    label=f_label,
                    data=file_bytes,
                    file_name=f_name,
                    mime=mime_type
                )
